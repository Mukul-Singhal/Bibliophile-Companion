import keyboard
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import ImageGrab
import clipboard
import time

last = None
def save_2_doc(what):
    global last
    document.add_heading(datetime.now().strftime("%H:%M:%S"),2)
    para = document.add_paragraph()
    if(what == 'img'):
        last.save('tmp.png','PNG')
        r = para.add_run()
        r.add_picture('tmp.png')
    elif(what == 'words'):
        document.add_paragraph(last)
    else:
        raise("NOT DEFINE")

def image_handler():
    global last
    if(str(last)[53:58] == str(ImageGrab.grabclipboard())[53:58]):
        print("Same Iamge")
        return
    
    print(ImageGrab.grabclipboard())
    last = ImageGrab.grabclipboard()
    save_2_doc('img')

def word_handler():
    global last
    if(last == clipboard.paste() or len(clipboard.paste().strip()) == 0):
        print("Same sentence")
        return

    print(clipboard.paste())
    last = clipboard.paste()
    save_2_doc('words')

if __name__ =="__main__":
    document = Document()
    document.add_heading(datetime.now().strftime("%d-%m-%Y"), 0)


    while(1):
        if keyboard.is_pressed('ctrl+C') or keyboard.is_pressed('prtscn') == True:
            time.sleep(1)
            if(ImageGrab.grabclipboard() != None):
                image_handler()
            elif (clipboard.paste() != None):
                word_handler()
            else:
                print("UNKNOWN KEY Pressed")
        if keyboard.is_pressed('esc+X') == True:
            break
    document.save('demo2.docx')
    document.add_page_break()

